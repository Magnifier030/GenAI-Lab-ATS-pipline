from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import defaultdict
from tqdm import tqdm

def hex_to_rgb(hex_color):
    """將十六進制顏色代碼轉換為 RGB 顏色"""
    hex_color = hex_color.lstrip('#').replace(';', '')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

color_map = {
    'yellow': WD_COLOR_INDEX.YELLOW,
    'green': WD_COLOR_INDEX.BRIGHT_GREEN,
    'red': WD_COLOR_INDEX.RED,
    'cyan': WD_COLOR_INDEX.TURQUOISE,
    'blue' : WD_COLOR_INDEX.BLUE,
    'darkBlue': WD_COLOR_INDEX.DARK_BLUE,
    'darkRed': WD_COLOR_INDEX.DARK_RED,
    'darkCyan': WD_COLOR_INDEX.TEAL,
    'darkMagenta': WD_COLOR_INDEX.VIOLET,
    'darkYellow': WD_COLOR_INDEX.DARK_YELLOW,
    'magenta' :WD_COLOR_INDEX.PINK,
    'lightGray' : WD_COLOR_INDEX.GRAY_25,
    'darkGray' : WD_COLOR_INDEX.GRAY_50,
    'darkGreen' : WD_COLOR_INDEX.GREEN,
    'white' : WD_COLOR_INDEX.WHITE
}

def apply_highlight(run, color):
    """將 highlight 顏色應用到 run"""
    if color in color_map:
        run.font.highlight_color = color_map[color]
    else:
        return None
        # rgb = hex_to_rgb(color)
        # run.font.highlight_color = RGBColor(rgb[0], rgb[1], rgb[2])

def apply_color(run, color_code):
    """將顏色應用到 run"""
    rgb = hex_to_rgb(color_code)
    
    run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
    
def apply_highlight_and_color(run, highlight_color, font_color):
    """將 highlight 和 font 顏色應用到 run"""
    if highlight_color in color_map:
        run.font.highlight_color = color_map[highlight_color]
    else:
        highlight_rgb = hex_to_rgb(highlight_color)
        run.font.highlight_color = RGBColor(highlight_rgb[0], highlight_rgb[1], highlight_rgb[2])
    
    if font_color in color_map:
        run.font.color.rgb = color_map[font_color]
    else:
        font_rgb = hex_to_rgb(font_color)
        run.font.color.rgb = RGBColor(font_rgb[0], font_rgb[1], font_rgb[2])

def set_table_borders(table):
    """為表格添加邊框"""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

def process_highlight(highlight, paragraph, changes):
    """處理 highlight 標籤並將文字和顏色應用到段落"""
    if highlight.get('value'):
        color = highlight.get('value')
    else:
        color = None

    for part in highlight.contents:
        if isinstance(part, str):
            run = paragraph.add_run(part)
            if color:
                apply_highlight(run, color)
                changes.append(part + " [顏色: " + color + "]")
        else:
            sub_run = paragraph.add_run(part.get_text(strip=True))
            if part.get('value'):
                apply_color(sub_run, part.get('value'))
                changes.append(part.get_text(strip=True) + " [顏色: " + part.get('value') + "]")
            elif color:
                apply_highlight(sub_run, color)
                changes.append(part.get_text(strip=True) + " [顏色: " + color + "]")

def process_tag(tag, doc, changes):
    if tag.name == "table":
        table_data = []
        for row in tag.find_all('row'):
            row_data = []
            for cell in row.find_all('column'):
                cell_text = []
                for cell_clip in cell.find_all(recursive=False):
                    if cell_clip.name == "highlight":
                        highlight_value = cell_clip.get('value')

                        for color in cell_clip.find_all('color'):
                            color_value = color.get('value')
                            color_text = color.get_text()
                            # print('color_text:', color_text)
                            # print('color_value:', color_value)
                            # print("color:",color)

                            cell_text.append(
                                [
                                    color_text,
                                    color_value,
                                    highlight_value
                                ]
                            )

                    else:
                        color_value = cell_clip.get('value')

                        for color in cell_clip.find_all('highlight'):
                            highlight_value = color.get('value')
                            highlight_text = color.get_text()
                            # print('highlight_text:', highlight_text)
                            # print('highlight_value:', highlight_value)

                            cell_text.append(
                                [
                                    highlight_text,
                                    color_value,
                                    highlight_value
                                ]
                            )
                row_data.append(cell_text)
                
        
            table_data.append(row_data)
            # print('table_data:', table_data)
        return table_data
        
                            
                            
    # else:
    #     p = doc.add_paragraph()
    #     combined_run = p.add_run()
    #     for child in tag.children:
    #         if child.name == 'highlight':
    #             run_text = []
    #             colors = []
    #             for part in child.contents:
    #                 if isinstance(part, str):
    #                     run_text.append(part)
    #                     colors.append(child.get('value'))
    #                 else:
    #                     run_text.append(part.get_text(strip=True))
    #                     colors.append(part.get('value') if part.get('value') else child.get('value'))
    #             combined_text = "".join(run_text)
    #             combined_run.text = combined_text
    #             for i, text in enumerate(run_text):
    #                 sub_run = p.add_run(text)
    #                 if colors[i]:
    #                     if colors[i] in color_map:
    #                         apply_highlight(sub_run, colors[i],)
    #                     else:
    #                         apply_color(sub_run, colors[i])
    #                     changes.append(text + " [顏色: " + colors[i] + "]")
    #         else:
    #             run = p.add_run(child.get_text(strip=True))
    #             if child.get('color'):
    #                 apply_color(run, child.get('color'))
    #                 changes.append(child.get_text(strip=True) + " [顏色: " + child.get('color') + "]")
                    # print('changes2:', changes)

def format_command_and_criteria(command_and_criteria):
    try:
        soup = BeautifulSoup(command_and_criteria, 'html.parser')
        return soup
    except Exception as e:
        return BeautifulSoup("", 'html.parser')

# 讀取 JSON 文件



def create_document(datas, output_file_path):
    testplan_dict = []
    i = 0
    for data in datas:
        i+=1
        command_and_criteria = data['instruction']
        soup = format_command_and_criteria(command_and_criteria)
        testplan_dict.append({
            'commandAndCriteria': soup,
            'test_item' : data['test_item'],
            'test_station' : data['test_station'],
            'generative_commands': data['generative_commands'],
            'generative_criterias' : data['generative_criterias']
        })

    # 建立 Word 文件
    doc = Document()
    content = []
    changes = []
    for item in testplan_dict:
        for tag in item['commandAndCriteria'].find_all(recursive=False):
            content.append(
                {
                    'test_station' : item['test_station'],
                    'test_item' : item['test_item'],
                    'commandAndCriteria' : process_tag(tag, doc, changes),
                    'generative_commands' : item['generative_commands'],
                    'generative_criterias' : item['generative_criterias']
                }
            )
    #key = test_station
    #value = test_items
    # items() = (key, value)
    i = 0
    word_table = doc.add_table(rows=len(content), cols=6)
    word_table.cell(i, 0).paragraphs[0].add_run('test_station')
    word_table.cell(i, 1).paragraphs[0].add_run('command')
    word_table.cell(i, 2).paragraphs[0].add_run('criteria')
    word_table.cell(i, 3).paragraphs[0].add_run('generative_command')
    word_table.cell(i, 4).paragraphs[0].add_run('generative_criteria')
    set_table_borders(word_table)
    for test_item_content in tqdm(content, desc="Processing items"):
        i += 1
            
        row_data = test_item_content['commandAndCriteria'][0]

        cell = word_table.cell(i, 0)
        run = cell.paragraphs[0].add_run(test_item_content['test_station'])
        cell = word_table.cell(i, 1)
        run = cell.paragraphs[0].add_run(test_item_content['test_item'])
        for j, cell_data in enumerate(row_data):
            cell = word_table.cell(i, j+2)
            print(cell_data)
            combined_run = cell.paragraphs[0].add_run()
            for text, color_value, highlight_value in cell_data:
                sub_run = cell.paragraphs[0].add_run(text)
                if highlight_value != "None" and highlight_value is not None:
                    apply_highlight(sub_run, highlight_value)
                if color_value != "None" and color_value is not None:
                    apply_color(sub_run, color_value)
        
        cell = word_table.cell(i, 4)
        run = cell.paragraphs[0].add_run("\n".join(test_item_content['generative_commands']))
        cell = word_table.cell(i, 5)
        run = cell.paragraphs[0].add_run("\n".join(test_item_content['generative_criterias']))


    # 保存 Word 文件
    doc.save(output_file_path)

    print("重建的 testplan 已保存為文件: " + output_file_path)